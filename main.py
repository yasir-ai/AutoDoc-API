# filename: main.py
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse, JSONResponse
from docx import Document
from docx.shared import RGBColor
import json
import os

app = FastAPI()

@app.post("/modify-docx")
async def modify_docx(file: UploadFile = File(...), edits: str = Form(...)):
    try:
        edits = json.loads(edits)
    except json.JSONDecodeError:
        return JSONResponse(content={"error": "Invalid JSON in 'edits'"}, status_code=400)

    # Set paths in a known writable directory (safe for Render/Railway)
    temp_dir = "/tmp"
    input_path = os.path.join(temp_dir, "uploaded.docx")
    output_path = os.path.join(temp_dir, "modified_output.docx")

    # Save uploaded file
    with open(input_path, "wb") as buffer:
        buffer.write(await file.read())

    # Modify the document
    doc = Document(input_path)

    replacements = edits.get("replacements", {})
    highlight_sections = edits.get("highlightSections", [])
    inline_notes = edits.get("inlineNotes", {})

    # Apply replacements
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, val)

    # Highlight specified sections
    for para in doc.paragraphs:
        if any(section in para.text for section in highlight_sections):
            run = para.add_run("  ← [highlighted section]")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 255)  # blue

    # Add inline notes
    for para in doc.paragraphs:
        for trigger, note in inline_notes.items():
            if trigger in para.text:
                para.text += f"  *{note}*"

    # Save output
    doc.save(output_path)

    # ✅ Log + check file existence
    if os.path.exists(output_path):
        print(f"[✅] Returning modified file from {output_path}")
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="Modified_Territory_Plan.docx"
        )
    else:
        print(f"[❌] File not found at {output_path}")
        return JSONResponse(content={"error": "File not found after saving"}, status_code=500)
